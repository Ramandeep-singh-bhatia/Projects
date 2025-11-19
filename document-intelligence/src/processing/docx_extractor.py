"""
DOCX document extractor.
Handles Microsoft Word documents with structure preservation.
"""

from pathlib import Path
from typing import Dict, Any
from docx import Document
from datetime import datetime
from src.processing.base_extractor import BaseExtractor, ExtractedDocument
from src.utils.logger import app_logger as logger


class DOCXExtractor(BaseExtractor):
    """Extractor for DOCX documents."""

    def can_handle(self, file_path: Path) -> bool:
        """Check if file is a DOCX."""
        return self._get_file_extension(file_path) == 'docx'

    def extract(self, file_path: Path) -> ExtractedDocument:
        """
        Extract content and metadata from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            ExtractedDocument with content and metadata
        """
        logger.info(f"Extracting DOCX: {file_path}")

        try:
            doc = Document(file_path)

            # Extract content with structure
            content = self._extract_content(doc)

            # Extract metadata
            metadata = self._extract_metadata(doc, file_path)

            # Count words
            word_count = self._count_words(content)
            language = self._detect_language(content)

            logger.info(f"Successfully extracted DOCX: {word_count} words")

            return ExtractedDocument(
                content=content,
                metadata=metadata,
                page_count=None,  # DOCX doesn't have fixed pages
                word_count=word_count,
                language=language,
                title=metadata.get('title'),
                author=metadata.get('author'),
                created_date=metadata.get('created_date'),
                modified_date=metadata.get('modified_date')
            )

        except Exception as e:
            logger.error(f"Failed to extract DOCX {file_path}: {str(e)}")
            raise

    def _extract_content(self, doc: Document) -> str:
        """Extract content while preserving document structure."""
        content = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Identify headings by style
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    content.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {text}\n")
                else:
                    content.append(text)

        # Extract tables
        for table_num, table in enumerate(doc.tables, 1):
            content.append(f"\n[Table {table_num}]")
            table_content = self._extract_table(table)
            content.append(table_content)

        return "\n".join(content)

    def _extract_table(self, table) -> str:
        """Extract table content."""
        table_text = []

        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            table_text.append(row_text)

        return "\n".join(table_text)

    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        metadata = {}

        try:
            # Core properties
            core_props = doc.core_properties

            metadata['title'] = core_props.title or ""
            metadata['author'] = core_props.author or ""
            metadata['subject'] = core_props.subject or ""
            metadata['keywords'] = core_props.keywords or ""
            metadata['category'] = core_props.category or ""
            metadata['comments'] = core_props.comments or ""

            # Dates
            if core_props.created:
                metadata['created_date'] = core_props.created

            if core_props.modified:
                metadata['modified_date'] = core_props.modified

            # Count statistics
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)

            # Get sections count
            metadata['section_count'] = len(doc.sections)

        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {str(e)}")

        return metadata
